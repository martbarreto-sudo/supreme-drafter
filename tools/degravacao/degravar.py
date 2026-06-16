#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Degravação de audiência via Gemini (multimodal).

Recebe um arquivo de áudio/vídeo e devolve a degravação em PT-BR, com blocos de
timestamp e IDENTIFICAÇÃO DOS DEPOENTES — no padrão usado pela banca. Saída em
.txt e, opcionalmente, .docx formatado.

Uso (na sua máquina, com o Gemini logado por API key):
    export GEMINI_API_KEY="sua_chave"   # https://aistudio.google.com/apikey
    python3 degravar.py audiencia.mp4 --docx \
        --depoentes "Juiz: Dr. Fulano; MP: Dr. Bruno; Defesa: Dr. Ayron; Vítima: Luan; Testemunha: Renan"

Dependências:
    pip install google-genai          # transcrição
    pip install python-docx           # apenas se usar --docx
"""
import argparse
import os
import re
import sys
import time

PROMPT_BASE = """Você é um degravador forense de audiências judiciais brasileiras.
Transcreva o áudio/vídeo a seguir em português do Brasil, com MÁXIMA fidelidade.

REGRAS:
1. Divida em blocos por trecho, cada um iniciado por um timestamp no formato
   "(M:SS - M:SS)" (início e fim do trecho), seguido do texto na linha de baixo.
2. IDENTIFIQUE OS DEPOENTES: prefixe cada fala com o papel/nome de quem fala
   (ex.: "Juiz:", "MP:", "Defesa:", "Vítima (Luan):", "Testemunha (Renan):").
   Use o elenco informado abaixo para nomear corretamente; se não for possível
   identificar com segurança, use "Interlocutor não identificado:".
3. Transcrição VERBATIM — não corrija gramática, não resuma, não interprete.
4. Trechos inaudíveis: marque como [inaudível] (ou [inaudível 00:00] se útil).
5. Não invente conteúdo. Se houver dúvida sobre uma palavra, use [?] após ela.
6. Não inclua comentários seus fora da transcrição.
"""

HEADER = """DEGRAVAÇÃO — {arquivo}
Gerado por Gemini ({modelo}) em {data}
{elenco}
{sep}
"""

TS_RE = re.compile(r"^\(\d{1,2}:\d{2}")            # início de bloco de timestamp
SPK_RE = re.compile(r"^([^:]{1,40}:)(.*)$")         # "Depoente: fala"


def transcribe(audio, modelo, depoentes):
    """Chama o Gemini e retorna o texto da degravação."""
    try:
        from google import genai
    except ImportError:
        sys.exit("ERRO: dependência ausente. Instale com:  pip install google-genai")

    client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
    print(f"↑ Enviando '{audio}' ...", file=sys.stderr)
    f = client.files.upload(file=audio)
    while getattr(f, "state", None) and str(f.state).endswith("PROCESSING"):
        time.sleep(3)
        f = client.files.get(name=f.name)
    if getattr(f, "state", None) and str(f.state).endswith("FAILED"):
        sys.exit("ERRO: o Gemini falhou ao processar o arquivo enviado.")

    prompt = PROMPT_BASE + (f"\nELENCO PARA IDENTIFICAÇÃO:\n{depoentes}\n" if depoentes else "")
    print(f"⚙  Degravando com {modelo} ...", file=sys.stderr)
    resp = client.models.generate_content(model=modelo, contents=[prompt, f])
    texto = (resp.text or "").strip()
    if not texto:
        sys.exit("ERRO: resposta vazia do modelo.")
    return texto


def write_txt(texto, path, meta):
    with open(path, "w", encoding="utf-8") as out:
        out.write(HEADER.format(**meta) + "\n" + texto + "\n")


def write_docx(texto, path, meta):
    """Gera .docx no layout da banca: timestamp em destaque + falas com depoente em negrito."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        sys.exit("ERRO: para --docx instale:  pip install python-docx")

    MINT = RGBColor(0x10, 0xB9, 0x81)
    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    title = doc.add_paragraph()
    r = title.add_run(f"DEGRAVAÇÃO — {meta['arquivo']}")
    r.bold = True
    r.font.size = Pt(13)
    sub = doc.add_paragraph()
    rs = sub.add_run(f"Gemini ({meta['modelo']}) · {meta['data']}\n{meta['elenco']}")
    rs.italic = True
    rs.font.size = Pt(8.5)
    rs.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()

    for raw in texto.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if TS_RE.match(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = MINT
        else:
            p = doc.add_paragraph()
            m = SPK_RE.match(line)
            if m:
                p.add_run(m.group(1)).bold = True
                p.add_run(m.group(2))
            else:
                p.add_run(line)
    doc.save(path)


def main():
    ap = argparse.ArgumentParser(description="Degravação de audiência via Gemini.")
    ap.add_argument("audio", help="arquivo de áudio/vídeo (mp3, m4a, wav, mp4, ...)")
    ap.add_argument("--saida", "-o", default=None, help="arquivo .txt de saída (padrão: <audio>.degravacao.txt)")
    ap.add_argument("--docx", action="store_true", help="também gerar .docx formatado")
    ap.add_argument("--depoentes", "-d", default="", help="elenco: 'Papel: Nome; Papel: Nome'")
    ap.add_argument("--modelo", "-m", default="gemini-2.5-pro",
                    help="modelo Gemini (padrão gemini-2.5-pro; use gemini-2.5-flash p/ rapidez/custo)")
    args = ap.parse_args()

    if not os.path.isfile(args.audio):
        sys.exit(f"ERRO: arquivo não encontrado: {args.audio}")
    if not os.environ.get("GEMINI_API_KEY"):
        sys.exit("ERRO: defina GEMINI_API_KEY (https://aistudio.google.com/apikey).")

    saida = args.saida or (os.path.splitext(args.audio)[0] + ".degravacao.txt")
    meta = {
        "arquivo": os.path.basename(args.audio),
        "modelo": args.modelo,
        "data": time.strftime("%Y-%m-%d %H:%M"),
        "elenco": f"Elenco informado: {args.depoentes}" if args.depoentes else "Elenco: não informado",
        "sep": "=" * 60,
    }

    texto = transcribe(args.audio, args.modelo, args.depoentes)
    write_txt(texto, saida, meta)
    print(f"✓ TXT  → {saida}", file=sys.stderr)
    if args.docx:
        dpath = os.path.splitext(saida)[0] + ".docx"
        write_docx(texto, dpath, meta)
        print(f"✓ DOCX → {dpath}", file=sys.stderr)


if __name__ == "__main__":
    main()
